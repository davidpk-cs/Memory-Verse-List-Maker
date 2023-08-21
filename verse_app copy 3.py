import tkinter as tk
from tkinter import ttk
from tkinter import messagebox

import pythonbible as bible

from docx import Document


BibleVersions = ["KJV", "NIV"]

userVerseList = []

versesOnDisplay = []

versesSearched = ""

def get_bible_book(number):
    bible_books = [
        "Genesis", "Exodus", "Leviticus", "Numbers", "Deuteronomy", "Joshua", "Judges",
        "Ruth", "1 Samuel", "2 Samuel", "1 Kings", "2 Kings", "1 Chronicles", "2 Chronicles",
        "Ezra", "Nehemiah", "Esther", "Job", "Psalms", "Proverbs", "Ecclesiastes", "Song of Solomon",
        "Isaiah", "Jeremiah", "Lamentations", "Ezekiel", "Daniel", "Hosea", "Joel", "Amos", "Obadiah",
        "Jonah", "Micah", "Nahum", "Habakkuk", "Zephaniah", "Haggai", "Zechariah", "Malachi",
        "Matthew", "Mark", "Luke", "John", "Acts", "Romans", "1 Corinthians", "2 Corinthians",
        "Galatians", "Ephesians", "Philippians", "Colossians", "1 Thessalonians", "2 Thessalonians",
        "1 Timothy", "2 Timothy", "Titus", "Philemon", "Hebrews", "James", "1 Peter", "2 Peter",
        "1 John", "2 John", "3 John", "Jude", "Revelation"
    ]
    
    if 1 <= number <= 66:
        return bible_books[number - 1]
    else:
        return "Invalid input. Please enter a number between 1 and 66."

if __name__ == "__main__":


    def referenceSearch():

        try:

            global versesOnDisplay, versesSearched

            userReference = referenceEntry.get().strip()
            '''
            for num in ["1", "2", "3", "1st", "2nd", "3rd"]:
                if(userReference.startswith(num)):
                    userReference = num + userReference
            '''

            verseList = bible.get_references(userReference)

            versesOnDisplay = bible.convert_references_to_verse_ids(verseList)

            versesSearched = ""

            for verses in versesOnDisplay:
                versesSearched += "\n" + bible.get_verse_text(verses, version=bible.Version.KING_JAMES)

            verseDisplaySection.delete("1.0", "end")
            verseDisplaySection.insert("1.0", versesSearched)



            


            
            reference = bible.NormalizedReference(bible.Book.JOHN, 3, 16, 3, 16)

            # Convert reference to verse ID and get verse text
            verse_ids = bible.convert_reference_to_verse_ids(reference)
            verse_text = bible.get_verse_text(verse_ids[0])  # Assuming there's only one verse ID
            print(verse_text)

        except:
            messagebox.showerror("Error", "Not a valid reference")
        

    def verseAdd():

        global userVerseList, versesOnDisplay
        
        userVerseList.extend(versesOnDisplay)

        fullPageDisplay.delete("1.0", "end")

        toDisplayInFull = ""

        

        for id in userVerseList:
            if(isinstance(id, int)):
                stringID = str(id)
                if(len(str(id)) == 7):
                    stringID = "0" + str(id)
                toDisplayInFull += get_bible_book(int(stringID[0:2])) + " " + str(int(stringID[2: 5])) + ":" + str(int(stringID[5:])) + " - " + bible.get_verse_text(id) + "\n\n"
            else:
                stringID = str(id[0])
                finalID = str(id[-1])
                if(len(str(finalID)) == 7):
                    finalID = "0" + str(finalID)
                if(len(str(stringID)) == 7):
                    stringID = "0" + str(id)
               
                toDisplayInFull += get_bible_book(int(stringID[0:2])) + " " + str(int(stringID[2:5])) + ":" + str(int(stringID[5:])) + "-" + str(int(finalID[5:])) + "\n"
                for subID in id:
                    stringID = str(subID)
                    if(len(str(stringID)) == 7):
                        stringID = "0" + str(subID)
                    toDisplayInFull += str(int(stringID[5:])) + "-" + bible.get_verse_text(subID) + "\n"

                toDisplayInFull += "\n\n"

        fullPageDisplay.insert("1.0", toDisplayInFull)

    def groupVerseAdd():
        
        global userVerseList, versesOnDisplay

        if(len(versesOnDisplay) == 1):
            verseAdd()
            return

        for i in range(len(versesOnDisplay) - 1):
            if(versesOnDisplay[i] != versesOnDisplay[i + 1] - 1):
                messagebox.showerror("Error", "Non-Consecutive Verses")
                return
        
        userVerseList.append(versesOnDisplay)

        fullPageDisplay.delete("1.0", "end")

        toDisplayInFull = ""

        

        for id in userVerseList:
            if(isinstance(id, int)):
                stringID = str(id)
                if(len(str(id)) == 7):
                    stringID = "0" + str(id)  
                toDisplayInFull += get_bible_book(int(stringID[0:2])) + " " + str(int(stringID[2: 5])) + ":" + str(int(stringID[5:])) + " - " + bible.get_verse_text(id) + "\n\n"
            else:
                stringID = str(id[0])
                finalID = str(id[-1])
                if(len(str(finalID)) == 7):
                    finalID = "0" + str(finalID)
                if(len(str(stringID)) == 7):
                    stringID = "0" + str(stringID)
                toDisplayInFull += get_bible_book(int(stringID[0:2])) + " " + str(int(stringID[2: 5])) + ":" + str(int(stringID[5:])) + "-" + str(int(finalID[5:])) + "\n"
                for subID in id:
                    stringID = str(subID)
                    if(len(str(stringID)) == 7):
                        stringID = "0" + str(subID) 
                    toDisplayInFull += str(int(stringID[5:])) + "-" + bible.get_verse_text(subID) + "\n"

                toDisplayInFull += "\n\n"

        fullPageDisplay.insert("1.0", toDisplayInFull)



    def export():
        
        challenge = Document()

        challenge.add_heading("Test Challenge", level = 1)

        for id in userVerseList:
            if(isinstance(id, int)):
                stringID = str(id)
                if(len(str(id)) == 7):
                    stringID = "0" + str(id)  
                paragraph = challenge.add_paragraph(get_bible_book(int(stringID[0:2])) + " " + str(int(stringID[2: 5])) + ":" + str(int(stringID[5:])) + " - " + bible.get_verse_text(id) + "\n")
            else:
                stringID = str(id[0])
                finalID = str(id[-1])
                if(len(str(finalID)) == 7):
                    finalID = "0" + str(finalID)
                if(len(str(stringID)) == 7):
                    stringID = "0" + str(stringID)
                paragraph = challenge.add_paragraph(get_bible_book(int(stringID[0:2])) + " " + str(int(stringID[2: 5])) + ":" + str(int(stringID[5:])) + "-" + str(int(finalID[5:])) + "\n")
                for subID in id:
                    stringID = str(subID)
                    if(len(str(stringID)) == 7):
                        stringID = "0" + str(subID) 
                    paragraph.add_run(str(int(stringID[5:]))).font.superscript = True
                    paragraph.add_run(bible.get_verse_text(subID))
                paragraph.add_run("\n")


        challenge.save("test1")

    mainWindow = tk.Tk()
    mainWindow.geometry('1000x600')
    mainWindow.title("Memory Verse Challenge Maker")
    mainWindow.configure(borderwidth = 20)

    mainWindow.grid_columnconfigure(0, weight=3)  
    mainWindow.grid_columnconfigure(1, weight=1) 


    leftSide = tk.Frame(mainWindow)

    rightSide = tk.Frame(mainWindow)
    rightSide.grid_columnconfigure(0, weight=1)
    rightSide.grid_rowconfigure(0, weight=1)

    queryHolder = tk.Frame(leftSide)


    #------#
    #leftSide of the window
    selectedVersion = tk.StringVar()
    versionSelector = ttk.OptionMenu(queryHolder, selectedVersion, BibleVersions[0], *BibleVersions)
    versionSelector.pack(side = "left")
    
    referenceEntry = tk.Entry(queryHolder)
    referenceEntry.pack(side = "left")


    verseFindButton = ttk.Button(queryHolder, text = "Search", command = referenceSearch)
    verseFindButton.pack(side = "left")


    queryHolder.pack(side = "top")

    verseDisplaySection = tk.Text(leftSide, width = 25)
    verseDisplaySection.pack(fill = "y")

    addHolder = ttk.Frame(leftSide)

    addVerseButton = ttk.Button(addHolder, text = "Add", command = verseAdd)
    addVerseButton.pack(side = "left")

    addVerseGroup = ttk.Button(addHolder, text = "Add Group", command = groupVerseAdd)
    addVerseGroup.pack(side = "right")

    addHolder.pack()

    leftSide.grid(row=0, column=0, sticky="nsew")
    #------#
    #end leftSide


    #------#
    #rightSide
    fullDisplay = tk.Frame(rightSide)
    fullDisplay.grid_columnconfigure(0, weight = 0)  
    fullDisplay.grid_columnconfigure(1, weight = 1) 
    fullDisplay.grid_rowconfigure(0, weight = 1)
    fullDisplay.grid_rowconfigure(1, weight = 0)

    

    fullPageDisplay = tk.Text(fullDisplay)

    # Create a Scrollbar widget and associate it with the Text widget
    scrollbar = tk.Scrollbar(fullDisplay, command=fullPageDisplay.yview)

    scrollbar.grid(row = 0, column = 1, sticky = "nsw")
    fullPageDisplay.grid(row = 0, column = 0, sticky = "nse")

    

    # Connect the Text widget to the Scrollbar
    fullPageDisplay.config(yscrollcommand=scrollbar.set)

    fullDisplay.grid(row = 0, column = 0, sticky = "nsew")
    

    exportButton = ttk.Button(rightSide, text = "Export", command = export)
    exportButton.grid(row = 1, column = 0)

    rightSide.grid(row=0, column=1, sticky="nsew")
    #-------#
    #end rightSide


    mainWindow.mainloop()



